#!/usr/bin/env python3
"""Generate an editable trademark evidence compilation DOCX from JSON.

The JSON source remains the structured working record. Original screenshots,
downloads, and official PDFs are never modified by this script.
"""

from __future__ import annotations

import argparse
import json
import sys
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None


PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_DXA = 80
CELL_MARGIN_BOTTOM_DXA = 80
CELL_MARGIN_START_DXA = 120
CELL_MARGIN_END_DXA = 120

# compact_reference_guide preset with a named CJK-safe font override.
# Hiragino Sans GB ships with macOS and is available to both Microsoft Word
# and the local LibreOffice QA pipeline.
FONT_LATIN = "Hiragino Sans GB"
FONT_EAST_ASIA = "Hiragino Sans GB"
COLOR_HEADING = "2E5D7B"
COLOR_HEADING_DARK = "1F4058"
COLOR_MUTED = "5B6570"
COLOR_HEADER_FILL = "E8EEF5"
COLOR_LIGHT_FILL = "F5F7F9"

REQUIRED_ITEM_FIELDS = (
    "number",
    "title",
    "source_type",
    "source_name",
    "query_date",
    "content",
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate an editable Word evidence compilation from JSON."
    )
    parser.add_argument("input_json", type=Path, help="UTF-8 evidence manifest JSON")
    parser.add_argument("output_docx", type=Path, help="Output .docx path")
    parser.add_argument(
        "--allow-placeholders",
        action="store_true",
        help="Allow template placeholders such as '待填写/待确认'.",
    )
    return parser.parse_args()


def load_manifest(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ValueError(f"输入文件不存在：{path}") from exc
    except json.JSONDecodeError as exc:
        raise ValueError(f"JSON 格式错误：第 {exc.lineno} 行第 {exc.colno} 列") from exc
    if not isinstance(data, dict):
        raise ValueError("JSON 顶层必须是对象。")
    return data


def validate_manifest(data: dict[str, Any], allow_placeholders: bool) -> None:
    matter = data.get("matter")
    items = data.get("items")
    if not isinstance(matter, dict):
        raise ValueError("缺少 matter 对象。")
    if not isinstance(items, list) or not items:
        raise ValueError("items 必须是至少包含一项证据的数组。")

    required_matter = ("client", "trademark", "jurisdiction", "procedure_stage")
    missing_matter = [
        key for key in required_matter if not str(matter.get(key, "")).strip()
    ]
    if missing_matter:
        raise ValueError("matter 缺少字段：" + "、".join(missing_matter))

    seen_numbers: set[str] = set()
    for index, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"第 {index} 项证据必须是对象。")
        missing = [
            key for key in REQUIRED_ITEM_FIELDS if not str(item.get(key, "")).strip()
        ]
        if missing:
            raise ValueError(f"第 {index} 项证据缺少字段：" + "、".join(missing))
        number = str(item["number"]).strip()
        if number in seen_numbers:
            raise ValueError(f"证据编号重复：{number}")
        seen_numbers.add(number)
        image_paths = item.get("image_paths", [])
        if image_paths is not None and not isinstance(image_paths, list):
            raise ValueError(f"证据 {number} 的 image_paths 必须是数组。")

    if not allow_placeholders:
        serialized = json.dumps(data, ensure_ascii=False)
        if "待填写" in serialized or "待确认" in serialized:
            raise ValueError(
                "输入仍包含“待填写/待确认”占位符；生成模板请加 --allow-placeholders。"
            )


def set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_DXA),
        ("bottom", CELL_MARGIN_BOTTOM_DXA),
        ("start", CELL_MARGIN_START_DXA),
        ("end", CELL_MARGIN_END_DXA),
    ):
        tag = f"w:{edge}"
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError("表格列宽总和必须等于 9360 DXA。")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_row_repeat_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "1")


def set_run_font(
    run: Any,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
) -> None:
    run.font.name = FONT_LATIN
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph_runs(paragraph: Any, **kwargs: Any) -> None:
    for run in paragraph.runs:
        set_run_font(run, **kwargs)


def configure_style(
    style: Any,
    *,
    size: float,
    color: str,
    bold: bool,
    before: float,
    after: float,
    line_spacing: float = 1.0,
) -> None:
    style.font.name = FONT_LATIN
    r_pr = style._element.get_or_add_rPr()
    fonts = r_pr.get_or_add_rFonts()
    fonts.set(qn("w:ascii"), FONT_LATIN)
    fonts.set(qn("w:hAnsi"), FONT_LATIN)
    fonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.keep_with_next = True


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    configure_style(
        normal,
        size=10.5,
        color="000000",
        bold=False,
        before=0,
        after=6,
        line_spacing=1.25,
    )
    normal.paragraph_format.keep_with_next = False

    title_style = doc.styles["Title"]
    configure_style(
        title_style,
        size=24,
        color=COLOR_HEADING_DARK,
        bold=True,
        before=0,
        after=8,
    )
    title_p_pr = title_style._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)
    configure_style(
        doc.styles["Subtitle"],
        size=11,
        color=COLOR_MUTED,
        bold=False,
        before=0,
        after=18,
    )
    configure_style(
        doc.styles["Heading 1"],
        size=16,
        color=COLOR_HEADING,
        bold=True,
        before=18,
        after=10,
    )
    configure_style(
        doc.styles["Heading 2"],
        size=13,
        color=COLOR_HEADING,
        bold=True,
        before=14,
        after=7,
    )
    configure_style(
        doc.styles["Heading 3"],
        size=12,
        color=COLOR_HEADING_DARK,
        bold=True,
        before=10,
        after=5,
    )


def configure_section(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_DXA / 1440)
    section.page_height = Inches(PAGE_HEIGHT_DXA / 1440)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("商标证据汇编｜可编辑主文件")
    set_run_font(run, size=8.5, color=COLOR_MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    run = footer_p.add_run("第 ")
    set_run_font(run, size=8.5, color=COLOR_MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    run2 = footer_p.add_run(" 页")
    set_run_font(run2, size=8.5, color=COLOR_MUTED)


def add_label_value_table(doc: Document, rows: list[tuple[str, str]]) -> Any:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "字段"
    header_cells[1].text = "内容"
    set_row_repeat_header(table.rows[0])
    for cell in header_cells:
        set_cell_shading(cell, COLOR_HEADER_FILL)
        style_paragraph_runs(cell.paragraphs[0], size=9, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        cell.paragraphs[0].paragraph_format.line_spacing = 1.1
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value or "未提供"
        set_cell_shading(cells[0], COLOR_LIGHT_FILL)
        style_paragraph_runs(cells[0].paragraphs[0], size=9.5, bold=True)
        style_paragraph_runs(cells[1].paragraphs[0], size=9.5)
        for cell in cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.15
    set_table_geometry(table, [1800, 7560])
    return table


def add_index_table(doc: Document, items: list[dict[str, Any]]) -> Any:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ("编号", "证据名称", "来源类型", "查询日期", "记录号/页码", "核验状态")
    set_row_repeat_header(table.rows[0])
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = text
        set_cell_shading(cell, COLOR_HEADER_FILL)
        style_paragraph_runs(cell.paragraphs[0], size=9, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for item in items:
        row = table.add_row().cells
        record_ref = " / ".join(
            part
            for part in (
                str(item.get("record_no", "")).strip(),
                str(item.get("page_ref", "")).strip(),
            )
            if part
        )
        values = (
            str(item.get("number", "")),
            str(item.get("title", "")),
            str(item.get("source_type", "")),
            str(item.get("query_date", "")),
            record_ref or "未提供",
            "待核验" if str(item.get("uncertainty", "")).strip() else "已记录",
        )
        for index, value in enumerate(values):
            row[index].text = value
            style_paragraph_runs(row[index].paragraphs[0], size=8.5)
            row[index].paragraphs[0].paragraph_format.space_after = Pt(0)
            row[index].paragraphs[0].paragraph_format.line_spacing = 1.1
            if index in (0, 2, 3, 4, 5):
                row[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_geometry(table, [780, 2460, 1260, 1260, 1680, 1920])
    return table


def add_field_paragraph(
    doc: Document,
    label: str,
    value: str,
    *,
    italic: bool = False,
    color: str | None = None,
) -> Any:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.2
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, size=10.5, bold=True, color=COLOR_HEADING_DARK)
    value_run = paragraph.add_run(value or "未提供")
    set_run_font(value_run, size=10.5, italic=italic, color=color)
    return paragraph


def resolve_image_path(raw_path: str, manifest_path: Path) -> Path:
    candidate = Path(raw_path).expanduser()
    if not candidate.is_absolute():
        candidate = (manifest_path.parent / candidate).resolve()
    return candidate


def image_size_inches(
    path: Path, max_width: float = 6.1, max_height: float = 7.4
) -> tuple[float, float]:
    if Image is None:
        return max_width, max_height
    with Image.open(path) as image:
        width_px, height_px = image.size
        width_in = width_px / 96
        height_in = height_px / 96
        scale = min(max_width / width_in, max_height / height_in, 1.0)
        return width_in * scale, height_in * scale


def add_images(doc: Document, item: dict[str, Any], manifest_path: Path) -> None:
    image_paths = item.get("image_paths") or []
    if not image_paths:
        return
    doc.add_heading("证据图像", level=2)
    for image_index, raw_path in enumerate(image_paths, start=1):
        image_path = resolve_image_path(str(raw_path), manifest_path)
        if not image_path.is_file():
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(f"[图像未找到] {image_path}")
            set_run_font(run, size=9.5, color="9B1C1C", bold=True)
            continue
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        width_in, height_in = image_size_inches(image_path)
        try:
            inline_shape = paragraph.add_run().add_picture(
                str(image_path), width=Inches(width_in), height=Inches(height_in)
            )
            inline_shape._inline.docPr.set(
                "descr", f"{item['number']} 证据图像：{image_path.name}"
            )
            inline_shape._inline.docPr.set("title", image_path.name)
        except Exception as exc:
            error_p = doc.add_paragraph()
            run = error_p.add_run(f"[图像嵌入失败] {image_path}：{exc}")
            set_run_font(run, size=9.5, color="9B1C1C", bold=True)
            continue
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(2)
        caption.paragraph_format.space_after = Pt(8)
        run = caption.add_run(
            f"图 {item['number']}-{image_index}｜原始文件：{image_path.name}"
        )
        set_run_font(run, size=8.5, italic=True, color=COLOR_MUTED)


def add_evidence_item(
    doc: Document, item: dict[str, Any], manifest_path: Path
) -> None:
    heading = doc.add_heading(f"{item['number']}｜{item['title']}", level=1)
    heading.paragraph_format.keep_with_next = True
    record_ref = " / ".join(
        part
        for part in (
            str(item.get("record_no", "")).strip(),
            str(item.get("page_ref", "")).strip(),
        )
        if part
    )
    add_label_value_table(
        doc,
        [
            ("来源类型", str(item.get("source_type", ""))),
            ("来源名称", str(item.get("source_name", ""))),
            ("查询日期", str(item.get("query_date", ""))),
            ("记录号 / 页码", record_ref),
            ("来源 URL", str(item.get("url", ""))),
            ("原始文件位置", str(item.get("file_path", ""))),
        ],
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    add_field_paragraph(doc, "证据内容", str(item.get("content", "")))
    add_field_paragraph(doc, "代理分析", str(item.get("analysis", "")))
    uncertainty = str(item.get("uncertainty", "")).strip()
    if uncertainty:
        add_field_paragraph(
            doc, "待核实事项", uncertainty, italic=True, color="7A5A00"
        )
    add_images(doc, item, manifest_path)


def build_document(data: dict[str, Any], manifest_path: Path) -> Document:
    matter = data["matter"]
    items = data["items"]
    doc = Document()
    configure_styles(doc)
    configure_section(doc)

    doc.core_properties.title = "商标证据汇编"
    doc.core_properties.subject = "可编辑商标证据主文件"
    doc.core_properties.comments = "本 DOCX 为证据汇编主文件；原始证据另行只读保留。"

    title = doc.add_paragraph(style="Title")
    title_run = title.add_run("商标证据汇编")
    set_run_font(title_run, size=24, bold=True, color=COLOR_HEADING_DARK)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle_run = subtitle.add_run(
        f"{matter.get('trademark', '')}｜{matter.get('jurisdiction', '')}｜"
        f"{matter.get('procedure_stage', '')}"
    )
    set_run_font(subtitle_run, size=11, color=COLOR_MUTED)

    jurisdiction_authority = " / ".join(
        part
        for part in (
            str(matter.get("jurisdiction", "")).strip(),
            str(matter.get("authority", "")).strip(),
        )
        if part
    )
    version_date = " / ".join(
        part
        for part in (
            str(matter.get("version", "")).strip(),
            str(matter.get("prepared_date", "")).strip() or date.today().isoformat(),
        )
        if part
    )
    add_label_value_table(
        doc,
        [
            ("客户", str(matter.get("client", ""))),
            ("商标", str(matter.get("trademark", ""))),
            ("法域 / 主管机关", jurisdiction_authority),
            ("程序阶段", str(matter.get("procedure_stage", ""))),
            ("官方文号", str(matter.get("official_number", ""))),
            ("版本 / 编制日期", version_date),
        ],
    )

    summary = str(data.get("summary", "")).strip()
    if summary:
        doc.add_heading("编制说明", level=1)
        paragraph = doc.add_paragraph(summary)
        style_paragraph_runs(paragraph, size=10.5)

    doc.add_heading("证据目录", level=1)
    add_index_table(doc, items)

    for item in items:
        add_evidence_item(doc, item, manifest_path)

    doc.add_heading("复核记录", level=1)
    add_label_value_table(
        doc,
        [
            ("内容复核", "待复核"),
            ("来源复核", "待复核"),
            ("编号与附件复核", "待复核"),
            ("版式与页码复核", "待复核"),
            ("复核人 / 日期", "待填写"),
        ],
    )
    return doc


def main() -> int:
    args = parse_args()
    if args.output_docx.suffix.lower() != ".docx":
        print("输出文件必须使用 .docx 扩展名。", file=sys.stderr)
        return 2
    try:
        data = load_manifest(args.input_json)
        validate_manifest(data, args.allow_placeholders)
        document = build_document(data, args.input_json.resolve())
        args.output_docx.parent.mkdir(parents=True, exist_ok=True)
        document.save(args.output_docx)
    except (ValueError, OSError) as exc:
        print(f"生成失败：{exc}", file=sys.stderr)
        return 2
    print(f"已生成：{args.output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
