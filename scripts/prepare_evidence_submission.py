#!/usr/bin/env python3
"""Prepare trademark evidence submission files from a compiled source PDF.

Outputs:
1. One editable DOCX containing the evidence explanation/cover pages.
2. One editable DOCX containing the signed evidence directory.
3. One PDF per evidence item, excluding the explanation pages.
4. A machine-readable audit report.

The script removes only pagination objects added by the compilation workflow.
Pagination that is intrinsic to a source brochure, webpage printout, certificate,
or official document is preserved as evidence content.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from pypdf import PdfReader, PdfWriter
from pypdf.generic import ContentStream, DictionaryObject


FONT = "Hiragino Sans GB"
BLACK = "000000"
MUTED = "6B6B6B"
A4_WIDTH_POINTS = 595.276
A4_HEIGHT_POINTS = 841.890
MAX_PDF_BYTES = 50 * 1024 * 1024
MAX_PDF_FILES = 15

# standard_business_brief with a named legal_evidence_cover override:
# A4; CJK-safe font; no running header or page number; repeated applicant footer;
# page border; centered legal cover hierarchy.
DOCX_STYLE_TOKENS = {
    "preset": "standard_business_brief",
    "override": "legal_evidence_cover",
    "page_size": "A4 portrait",
    "margins_mm": {"top": 18, "right": 25, "bottom": 23, "left": 25},
    "header_distance_mm": 8,
    "footer_distance_mm": 13,
    "base_font": FONT,
    "base_size_pt": 12,
    "body_after_pt": 5,
    "body_line_spacing": 1.25,
    "page_border": {"style": "single", "size_eighth_points": 8, "space_pt": 24},
    "running_header": None,
    "running_footer": "applicant_only",
    "page_number": None,
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate editable evidence explanations and split evidence PDFs."
    )
    parser.add_argument("manifest", type=Path, help="UTF-8 JSON manifest")
    parser.add_argument("output_dir", type=Path, help="Final delivery directory")
    parser.add_argument(
        "--audit-report",
        type=Path,
        help="Audit JSON path; defaults to <output_dir>/制作校验记录.json",
    )
    return parser.parse_args()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_manifest(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise ValueError(f"清单不存在：{path}") from exc
    except json.JSONDecodeError as exc:
        raise ValueError(f"清单 JSON 错误：第 {exc.lineno} 行第 {exc.colno} 列") from exc
    if not isinstance(data, dict):
        raise ValueError("清单顶层必须是对象。")
    return data


def resolve_source_pdf(manifest_path: Path, data: dict[str, Any]) -> Path:
    raw = str(data.get("source_pdf", "")).strip()
    if not raw:
        raise ValueError("清单缺少 source_pdf。")
    candidate = Path(raw)
    if not candidate.is_absolute():
        candidate = (manifest_path.parent / candidate).resolve()
    if not candidate.is_file():
        raise ValueError(f"源 PDF 不存在：{candidate}")
    return candidate


def resolve_optional_reference(
    manifest_path: Path,
    data: dict[str, Any],
    path_field: str,
    hash_field: str,
) -> tuple[Path | None, str | None]:
    raw = str(data.get(path_field, "")).strip()
    if not raw:
        return None, None
    candidate = Path(raw)
    if not candidate.is_absolute():
        candidate = (manifest_path.parent / candidate).resolve()
    if not candidate.is_file():
        raise ValueError(f"参考文件不存在：{candidate}")
    actual_hash = sha256_file(candidate)
    expected_hash = str(data.get(hash_field, "")).strip().lower()
    if expected_hash and actual_hash.lower() != expected_hash:
        raise ValueError(
            f"参考文件 {candidate.name} 的 SHA-256 与清单不一致，已停止。"
        )
    return candidate, actual_hash


def load_supplemental_sources(
    manifest_path: Path,
    data: dict[str, Any],
) -> dict[str, dict[str, Any]]:
    raw_sources = data.get("supplemental_sources", {})
    if not isinstance(raw_sources, dict):
        raise ValueError("supplemental_sources 必须是对象。")
    sources: dict[str, dict[str, Any]] = {}
    for source_id, raw_specification in raw_sources.items():
        if not isinstance(raw_specification, dict):
            raise ValueError(f"补充来源 {source_id} 必须是对象。")
        raw_path = str(raw_specification.get("path", "")).strip()
        if not raw_path:
            raise ValueError(f"补充来源 {source_id} 缺少 path。")
        path = Path(raw_path)
        if not path.is_absolute():
            path = (manifest_path.parent / path).resolve()
        if not path.is_file():
            raise ValueError(f"补充来源不存在：{path}")
        actual_hash = sha256_file(path)
        expected_hash = str(raw_specification.get("sha256", "")).strip().lower()
        if expected_hash and actual_hash.lower() != expected_hash:
            raise ValueError(f"补充来源 {source_id} 的 SHA-256 与清单不一致。")
        reader = PdfReader(path)
        if reader.is_encrypted:
            raise ValueError(f"补充来源已加密：{path.name}")
        declared_page_count = int(raw_specification.get("page_count", len(reader.pages)))
        if declared_page_count != len(reader.pages):
            raise ValueError(
                f"补充来源 {source_id} 声明 {declared_page_count} 页，"
                f"实际 {len(reader.pages)} 页。"
            )
        sources[str(source_id)] = {
            "path": path,
            "sha256": actual_hash,
            "page_count": len(reader.pages),
            "reader": reader,
        }
    return sources


def validate_manifest(
    data: dict[str, Any],
    page_count: int,
    supplemental_sources: dict[str, dict[str, Any]],
) -> None:
    matter = data.get("matter")
    items = data.get("items")
    if not isinstance(matter, dict):
        raise ValueError("清单缺少 matter 对象。")
    required_matter = (
        "client",
        "trademark",
        "jurisdiction",
        "procedure_stage",
        "disputed_mark",
        "evidence_category",
        "applicant",
    )
    missing_matter = [
        field for field in required_matter if not str(matter.get(field, "")).strip()
    ]
    if missing_matter:
        raise ValueError("matter 缺少字段：" + "、".join(missing_matter))

    if not isinstance(items, list) or not items:
        raise ValueError("items 必须是非空数组。")
    if len(items) > MAX_PDF_FILES:
        raise ValueError(f"证据 PDF 数量超过指引上限 {MAX_PDF_FILES} 个。")

    expected_number = 1
    expected_logical_start = 1
    used_source_pages: set[int] = set()
    cover_pages: set[int] = set()
    for position, item in enumerate(items, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"第 {position} 项必须是对象。")
        for field in (
            "number",
            "cover_source_page",
            "evidence_source_pages",
            "logical_pages",
            "title",
            "source",
            "fact_to_prove",
            "materials",
        ):
            if field not in item:
                raise ValueError(f"第 {position} 项缺少字段：{field}")

        number = int(item["number"])
        if number != expected_number:
            raise ValueError(
                f"证据编号应连续：位置 {position} 应为 {expected_number}，实际为 {number}。"
            )
        expected_number += 1

        cover_page = int(item["cover_source_page"])
        if not 1 <= cover_page <= page_count:
            raise ValueError(f"证据{number}说明页超出源 PDF：{cover_page}")
        if cover_page in cover_pages:
            raise ValueError(f"说明页重复：{cover_page}")
        cover_pages.add(cover_page)

        source_range = item["evidence_source_pages"]
        logical_range = item["logical_pages"]
        if (
            not isinstance(source_range, list)
            or len(source_range) != 2
            or not isinstance(logical_range, list)
            or len(logical_range) != 2
        ):
            raise ValueError(f"证据{number}页码范围必须为两个整数。")
        source_start, source_end = map(int, source_range)
        logical_start, logical_end = map(int, logical_range)
        if not (1 <= source_start <= source_end <= page_count):
            raise ValueError(f"证据{number}源页范围无效：{source_range}")
        if logical_start != expected_logical_start:
            raise ValueError(
                f"证据{number}逻辑起始页应为 {expected_logical_start}，实际为 {logical_start}。"
            )
        source_pages = source_end - source_start + 1
        supplements_before = item.get("supplements_before", [])
        if not isinstance(supplements_before, list):
            raise ValueError(f"证据{number} supplements_before 必须是数组。")
        supplemental_page_count = 0
        for supplement in supplements_before:
            if not isinstance(supplement, dict):
                raise ValueError(f"证据{number}补充来源项必须是对象。")
            source_id = str(supplement.get("source_id", "")).strip()
            if source_id not in supplemental_sources:
                raise ValueError(f"证据{number}引用未知补充来源：{source_id}")
            pages = supplement.get("pages")
            if not isinstance(pages, list) or len(pages) != 2:
                raise ValueError(f"证据{number}补充来源页码范围必须为两个整数。")
            supplement_start, supplement_end = map(int, pages)
            max_page = int(supplemental_sources[source_id]["page_count"])
            if not (1 <= supplement_start <= supplement_end <= max_page):
                raise ValueError(
                    f"证据{number}补充来源 {source_id} 页码范围无效：{pages}"
                )
            supplemental_page_count += supplement_end - supplement_start + 1
        logical_pages = logical_end - logical_start + 1
        total_item_pages = source_pages + supplemental_page_count
        if total_item_pages != logical_pages:
            raise ValueError(
                f"证据{number}正文总页数 {total_item_pages} "
                f"与逻辑页数 {logical_pages} 不一致。"
            )
        expected_logical_start = logical_end + 1
        for source_page in range(source_start, source_end + 1):
            if source_page in used_source_pages:
                raise ValueError(f"证据正文源页重复：{source_page}")
            used_source_pages.add(source_page)

        if not str(item["title"]).strip():
            raise ValueError(f"证据{number}标题为空。")
        if "/" in str(item["title"]) or ":" in str(item["title"]):
            raise ValueError(f"证据{number}标题含文件名禁用字符 / 或 :。")
        if not isinstance(item["materials"], list):
            raise ValueError(f"证据{number} materials 必须是数组。")

    collision = used_source_pages & cover_pages
    if collision:
        raise ValueError("说明页与证据正文页重叠：" + "、".join(map(str, sorted(collision))))
    if len(used_source_pages) + len(cover_pages) != page_count:
        missing = sorted(set(range(1, page_count + 1)) - used_source_pages - cover_pages)
        raise ValueError("源 PDF 存在未归类页面：" + "、".join(map(str, missing)))


def is_global_compilation_pagination(operands: list[Any], operator: bytes) -> bool:
    if operator != b"BDC" or len(operands) < 2:
        return False
    tag, properties = operands[0], operands[1]
    if str(tag) != "/Artifact" or not isinstance(properties, DictionaryObject):
        return False
    return (
        str(properties.get("/Subtype", "")) == "/Header"
        and str(properties.get("/Type", "")) == "/Pagination"
        and properties.get("/Contents") is not None
    )


def is_compilation_logical_font(operands: list[Any], operator: bytes) -> bool:
    """Identify the compiler's standalone logical-page-number font."""
    if operator != b"Tf" or len(operands) < 2:
        return False
    font_name = str(operands[0])
    try:
        font_size = float(operands[1])
    except (TypeError, ValueError):
        return False
    return bool(re.fullmatch(r"/C\d+", font_name)) and abs(font_size - 1.0) < 0.001


def strip_compilation_pagination(page: Any, pdf: PdfWriter) -> tuple[int, int]:
    """Remove added logical (/C1, /C2, ...) and physical pagination."""
    content = ContentStream(page.get_contents(), pdf)
    operations = content.operations
    remove_indices: set[int] = set()

    logical_removed = 0
    index = 0
    while index < len(operations):
        if operations[index][1] != b"BT":
            index += 1
            continue
        end = index + 1
        while end < len(operations) and operations[end][1] != b"ET":
            end += 1
        if end >= len(operations):
            raise ValueError("发现未闭合的 PDF 文本对象，无法安全移除汇编逻辑页码。")
        uses_compilation_font = any(
            is_compilation_logical_font(operands, op)
            for operands, op in operations[index + 1 : end]
        )
        if uses_compilation_font:
            remove_indices.update(range(index, end + 1))
            logical_removed += 1
        index = end + 1

    physical_removed = 0
    index = 0
    while index < len(operations):
        operands, operator = operations[index]
        if not is_global_compilation_pagination(operands, operator):
            index += 1
            continue
        depth = 1
        end = index + 1
        while end < len(operations) and depth:
            current_operator = operations[end][1]
            if current_operator in (b"BDC", b"BMC"):
                depth += 1
            elif current_operator == b"EMC":
                depth -= 1
            end += 1
        if depth:
            raise ValueError("发现未闭合的 PDF 标记内容，无法安全移除汇编物理页码。")
        remove_indices.update(range(index, end))
        physical_removed += 1
        index = end

    content.operations = [
        operation
        for position, operation in enumerate(operations)
        if position not in remove_indices
    ]
    page.replace_contents(content)
    return logical_removed, physical_removed


def safe_evidence_filename(item: dict[str, Any]) -> str:
    logical_start, logical_end = map(int, item["logical_pages"])
    raw = (
        f"{int(item['number'])}.{str(item['title']).strip()}"
        f"（第{logical_start}—{logical_end}页）.pdf"
    )
    return re.sub(r'[\\\\/:*?"<>|]', "＿", raw)


def is_a4(width: float, height: float) -> bool:
    portrait = (
        abs(width - A4_WIDTH_POINTS) <= 1
        and abs(height - A4_HEIGHT_POINTS) <= 1
    )
    landscape = (
        abs(width - A4_HEIGHT_POINTS) <= 1
        and abs(height - A4_WIDTH_POINTS) <= 1
    )
    return portrait or landscape


def add_page_border(section: Any) -> None:
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgBorders"))
    if existing is not None:
        sect_pr.remove(existing)
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "24")
        border.set(qn("w:color"), "333333")
        pg_borders.append(border)
    sect_pr.append(pg_borders)


def set_run_font(
    run: Any,
    *,
    size: float,
    bold: bool = False,
    color: str = BLACK,
    character_spacing_twips: int | None = None,
) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    if character_spacing_twips is not None:
        spacing = r_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            r_pr.append(spacing)
        spacing.set(qn("w:val"), str(character_spacing_twips))


def configure_docx(data: dict[str, Any]) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.right_margin = Mm(25)
    section.bottom_margin = Mm(23)
    section.left_margin = Mm(25)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(13)
    add_page_border(section)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal_r_pr = normal._element.get_or_add_rPr()
    normal_fonts = normal_r_pr.get_or_add_rFonts()
    normal_fonts.set(qn("w:ascii"), FONT)
    normal_fonts.set(qn("w:hAnsi"), FONT)
    normal_fonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.clear()

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    applicant = str(data["matter"]["applicant"]).strip()
    run = footer.add_run(f"申请人：{applicant}")
    set_run_font(run, size=12)

    document.core_properties.title = "证据说明——第一类申请人主体、在先权利及品牌影响"
    document.core_properties.subject = str(data["matter"]["procedure_stage"])
    document.core_properties.author = applicant
    document.core_properties.keywords = "商标,无效宣告,证据说明,可编辑Word"
    return document


def add_spacer(document: Document, points: float) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.line_spacing = 0.1
    run = paragraph.add_run("\u00a0")
    set_run_font(run, size=1)


def add_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "333333")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_centered(
    document: Document,
    text: str,
    *,
    size: float,
    bold: bool = False,
    color: str = BLACK,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.0,
    character_spacing_twips: int | None = None,
    keep_with_next: bool = False,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    set_run_font(
        run,
        size=size,
        bold=bold,
        color=color,
        character_spacing_twips=character_spacing_twips,
    )
    return paragraph


def add_label_value(
    document: Document,
    label: str,
    value: str,
    *,
    after: float = 4,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Mm(7)
    paragraph.paragraph_format.right_indent = Mm(7)
    paragraph.paragraph_format.first_line_indent = Mm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.35
    label_run = paragraph.add_run(f"{label}：")
    set_run_font(label_run, size=12.2, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=12.2)
    return paragraph


def add_evidence_cover(
    document: Document,
    matter: dict[str, Any],
    item: dict[str, Any],
) -> None:
    number = int(item["number"])
    title = str(item["title"]).strip()
    long_cover = number in (6, 7)

    add_spacer(document, 10)
    add_centered(
        document,
        "商标评审案件证据材料",
        size=25,
        character_spacing_twips=80,
        after=26,
        keep_with_next=True,
    )
    add_rule(document)
    add_centered(
        document,
        f"证据{number}",
        size=30,
        bold=True,
        before=26,
        after=14 if long_cover else 20,
        keep_with_next=True,
    )
    title_size = 18.2 if number == 7 else (19.3 if number == 6 else 21.5)
    add_centered(
        document,
        title,
        size=title_size,
        bold=True,
        after=26 if long_cover else 46,
        line_spacing=1.15,
        keep_with_next=True,
    )
    add_centered(
        document,
        str(matter["evidence_category"]),
        size=11.5,
        color=MUTED,
        after=25 if long_cover else 44,
        keep_with_next=True,
    )
    add_label_value(document, "争议商标", str(matter["disputed_mark"]))
    add_label_value(document, "证据来源", str(item["source"]))
    add_label_value(
        document,
        "待证事实",
        str(item["fact_to_prove"]),
        after=7 if item["materials"] else 0,
    )
    materials = item["materials"]
    if materials:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Mm(7)
        paragraph.paragraph_format.right_indent = Mm(7)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run("本组材料：")
        set_run_font(run, size=11.8, bold=True)
        for position, material in enumerate(materials, start=1):
            material_paragraph = document.add_paragraph()
            material_paragraph.paragraph_format.left_indent = Mm(12)
            material_paragraph.paragraph_format.right_indent = Mm(6)
            material_paragraph.paragraph_format.first_line_indent = Mm(-5)
            material_paragraph.paragraph_format.space_before = Pt(0)
            material_paragraph.paragraph_format.space_after = Pt(1.5)
            material_paragraph.paragraph_format.line_spacing = 1.15
            run = material_paragraph.add_run(f"{position}. {material}")
            set_run_font(run, size=11.2)


def create_explanation_docx(data: dict[str, Any], output_path: Path) -> None:
    document = configure_docx(data)
    items = data["items"]
    for position, item in enumerate(items):
        add_evidence_cover(document, data["matter"], item)
        if position != len(items) - 1:
            document.add_page_break()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def set_cell_width_dxa(cell: Any, width_dxa: int) -> None:
    cell.width = Mm(width_dxa / 1440 * 25.4)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 35), ("bottom", 35), ("start", 80), ("end", 80)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def configure_table_geometry(table: Any, widths_dxa: list[int]) -> None:
    total_width = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width_dxa(cell, widths_dxa[index])


def set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        tr_pr.append(node)
    node.set(qn("w:val"), "1")


def set_row_cant_split(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:cantSplit"))
    if node is None:
        node = OxmlElement("w:cantSplit")
        tr_pr.append(node)
    node.set(qn("w:val"), "1")


def set_cell_fill(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_text(
    cell: Any,
    text: str,
    *,
    size: float = 9.2,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def configure_catalog_docx(data: dict[str, Any]) -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(14)
    section.right_margin = Mm(25)
    section.bottom_margin = Mm(14)
    section.left_margin = Mm(25)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(10)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal_r_pr = normal._element.get_or_add_rPr()
    normal_fonts = normal_r_pr.get_or_add_rFonts()
    normal_fonts.set(qn("w:ascii"), FONT)
    normal_fonts.set(qn("w:hAnsi"), FONT)
    normal_fonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    document.core_properties.title = "证据目录——第一类申请人主体、在先权利及品牌影响"
    document.core_properties.subject = str(data["matter"]["procedure_stage"])
    document.core_properties.author = str(data["matter"]["applicant"])
    document.core_properties.keywords = "商标,无效宣告,证据目录,签章"
    return document


def create_evidence_catalog_docx(
    data: dict[str, Any],
    output_path: Path,
) -> None:
    document = configure_catalog_docx(data)
    add_centered(document, "证据目录", size=23, bold=True, after=4)
    add_centered(
        document,
        str(data["matter"]["evidence_category"]),
        size=12.5,
        color=MUTED,
        after=6,
    )
    add_label_value(document, "程序阶段", str(data["matter"]["procedure_stage"]), after=0)
    add_label_value(document, "争议商标", str(data["matter"]["disputed_mark"]), after=0)
    add_label_value(document, "申请人", str(data["matter"]["applicant"]), after=4)

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    # 160 mm content width / about 9071 DXA.
    headers = ["序号", "证据名称", "证据来源", "待证事实", "页码", "件数", "备注"]
    header_row = table.rows[0]
    for index, header in enumerate(headers):
        set_cell_text(
            header_row.cells[index],
            header,
            size=9.3,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        set_cell_fill(header_row.cells[index], "E8EEF5")
    set_repeat_table_header(header_row)

    for item in data["items"]:
        row = table.add_row()
        logical_start, logical_end = map(int, item["logical_pages"])
        materials = item.get("materials", [])
        values = [
            str(item["number"]),
            str(item["title"]),
            str(item["source"]),
            str(item["fact_to_prove"]),
            f"第{logical_start}—{logical_end}页",
            "1件",
            f"含{len(materials)}组材料" if materials else "",
        ]
        for index, value in enumerate(values):
            set_cell_text(
                row.cells[index],
                value,
                size=8.1 if index in (1, 3) else 8.4,
                align=(
                    WD_ALIGN_PARAGRAPH.CENTER
                    if index in (0, 4, 5, 6)
                    else WD_ALIGN_PARAGRAPH.LEFT
                ),
            )
        set_row_cant_split(row)
    configure_table_geometry(
        table,
        [650, 2000, 1230, 2450, 900, 620, 1221],
    )

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(
        "声明：上述证据文件名称、顺序和页码范围应与电子提交文件保持一致。"
    )
    set_run_font(run, size=9)

    signature = document.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature.paragraph_format.space_before = Pt(4)
    signature.paragraph_format.space_after = Pt(0)
    run = signature.add_run(
        "申请人（签章）：________________    日期：________年____月____日"
    )
    set_run_font(run, size=9.5)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def create_evidence_pdfs(
    reader: PdfReader,
    data: dict[str, Any],
    output_dir: Path,
    supplemental_sources: dict[str, dict[str, Any]],
) -> list[dict[str, Any]]:
    records: list[dict[str, Any]] = []
    total_logical_removed = 0
    total_physical_removed = 0
    for item in data["items"]:
        writer = PdfWriter()
        source_start, source_end = map(int, item["evidence_source_pages"])
        item_logical_removed = 0
        item_physical_removed = 0
        supplemental_records: list[dict[str, Any]] = []
        for supplement in item.get("supplements_before", []):
            source_id = str(supplement["source_id"])
            supplement_start, supplement_end = map(int, supplement["pages"])
            supplemental_reader = supplemental_sources[source_id]["reader"]
            supplemental_logical_removed = 0
            supplemental_physical_removed = 0
            for supplement_page in range(supplement_start, supplement_end + 1):
                writer.add_page(supplemental_reader.pages[supplement_page - 1])
                output_page = writer.pages[-1]
                logical_removed, physical_removed = strip_compilation_pagination(
                    output_page, writer
                )
                supplemental_logical_removed += logical_removed
                supplemental_physical_removed += physical_removed
                width = float(output_page.mediabox.width)
                height = float(output_page.mediabox.height)
                if not is_a4(width, height):
                    raise ValueError(
                        f"补充来源 {source_id} 第{supplement_page}页不是 A4："
                        f"{width:.3f}×{height:.3f} pt"
                    )
            item_logical_removed += supplemental_logical_removed
            item_physical_removed += supplemental_physical_removed
            supplemental_records.append(
                {
                    "source_id": source_id,
                    "path": str(supplemental_sources[source_id]["path"]),
                    "sha256": str(supplemental_sources[source_id]["sha256"]),
                    "pages": [supplement_start, supplement_end],
                    "page_count": supplement_end - supplement_start + 1,
                    "logical_pagination_objects_removed": supplemental_logical_removed,
                    "physical_pagination_objects_removed": supplemental_physical_removed,
                }
            )
        for source_page in range(source_start, source_end + 1):
            writer.add_page(reader.pages[source_page - 1])
            output_page = writer.pages[-1]
            logical_removed, physical_removed = strip_compilation_pagination(
                output_page, writer
            )
            item_logical_removed += logical_removed
            item_physical_removed += physical_removed
            width = float(output_page.mediabox.width)
            height = float(output_page.mediabox.height)
            if not is_a4(width, height):
                raise ValueError(
                    f"源第{source_page}页不是 A4：{width:.3f}×{height:.3f} pt"
                )

        expected_main_pages = source_end - source_start + 1
        expected_pages = expected_main_pages + sum(
            int(record["page_count"]) for record in supplemental_records
        )
        main_physical_removed = item_physical_removed - sum(
            int(record["physical_pagination_objects_removed"])
            for record in supplemental_records
        )
        if main_physical_removed != expected_main_pages:
            raise ValueError(
                f"证据{item['number']}主汇编物理页码移除数 {main_physical_removed} "
                f"与主汇编页数 {expected_main_pages} 不一致。"
            )
        filename = safe_evidence_filename(item)
        output_path = output_dir / filename
        writer.add_metadata(
            {
                "/Title": str(item["title"]),
                "/Subject": str(data["matter"]["procedure_stage"]),
                "/Author": str(data["matter"]["applicant"]),
                "/Keywords": "商标证据,电子证据提交,无新增页码",
            }
        )
        with output_path.open("wb") as stream:
            writer.write(stream)
        size_bytes = output_path.stat().st_size
        if size_bytes > MAX_PDF_BYTES:
            raise ValueError(
                f"{filename} 为 {size_bytes / 1024 / 1024:.2f} MB，超过 50 MB。"
            )
        total_logical_removed += item_logical_removed
        total_physical_removed += item_physical_removed
        records.append(
            {
                "number": int(item["number"]),
                "filename": filename,
                "source_pages": [source_start, source_end],
                "supplemental_sources": supplemental_records,
                "logical_pages": list(map(int, item["logical_pages"])),
                "page_count": expected_pages,
                "logical_pagination_objects_removed": item_logical_removed,
                "physical_pagination_objects_removed": item_physical_removed,
                "size_bytes": size_bytes,
                "sha256": sha256_file(output_path),
            }
        )

    return records


def validate_output_pdfs(output_dir: Path, records: list[dict[str, Any]]) -> None:
    for record in records:
        path = output_dir / record["filename"]
        reader = PdfReader(path)
        if reader.is_encrypted:
            raise ValueError(f"输出 PDF 不应加密：{path.name}")
        if len(reader.pages) != int(record["page_count"]):
            raise ValueError(f"输出 PDF 页数不符：{path.name}")
        for page_number, page in enumerate(reader.pages, start=1):
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            if not is_a4(width, height):
                raise ValueError(f"{path.name} 第{page_number}页不是 A4。")
            content = ContentStream(page.get_contents(), reader)
            if any(
                is_global_compilation_pagination(operands, operator)
                for operands, operator in content.operations
            ):
                raise ValueError(f"{path.name} 第{page_number}页仍含汇编物理页码对象。")
            if any(
                is_compilation_logical_font(operands, operator)
                for operands, operator in content.operations
            ):
                raise ValueError(f"{path.name} 第{page_number}页仍含汇编逻辑页码对象。")


def main() -> int:
    args = parse_args()
    manifest_path = args.manifest.resolve()
    data = load_manifest(manifest_path)
    source_pdf = resolve_source_pdf(manifest_path, data)
    source_hash = sha256_file(source_pdf)
    expected_hash = str(data.get("source_sha256", "")).strip().lower()
    if expected_hash and source_hash.lower() != expected_hash:
        raise ValueError(
            "源 PDF SHA-256 与清单不一致，已停止，防止案件材料混用。"
        )
    guide_pdf, guide_hash = resolve_optional_reference(
        manifest_path,
        data,
        "guide_pdf",
        "guide_sha256",
    )

    reader = PdfReader(source_pdf)
    if reader.is_encrypted:
        raise ValueError("源 PDF 已加密，无法处理。")
    supplemental_sources = load_supplemental_sources(manifest_path, data)
    validate_manifest(data, len(reader.pages), supplemental_sources)

    output_dir = args.output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    docx_name = str(data.get("output_docx_name", "")).strip()
    if not docx_name.endswith(".docx"):
        raise ValueError("output_docx_name 必须以 .docx 结尾。")
    docx_path = output_dir / docx_name
    catalog_docx_name = str(data.get("output_catalog_docx_name", "")).strip()
    if not catalog_docx_name.endswith(".docx"):
        raise ValueError("output_catalog_docx_name 必须以 .docx 结尾。")
    catalog_docx_path = output_dir / catalog_docx_name

    create_explanation_docx(data, docx_path)
    create_evidence_catalog_docx(data, catalog_docx_path)
    records = create_evidence_pdfs(
        reader,
        data,
        output_dir,
        supplemental_sources,
    )
    validate_output_pdfs(output_dir, records)

    audit_path = (
        args.audit_report.resolve()
        if args.audit_report
        else output_dir / "制作校验记录.json"
    )
    audit_path.parent.mkdir(parents=True, exist_ok=True)
    audit_data = {
        "checked_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "manifest": str(manifest_path),
        "manifest_sha256": sha256_file(manifest_path),
        "source_pdf": str(source_pdf),
        "source_sha256": source_hash,
        "source_page_count": len(reader.pages),
        "guide_pdf": str(guide_pdf) if guide_pdf else None,
        "guide_sha256": guide_hash,
        "supplemental_sources": {
            source_id: {
                "path": str(source["path"]),
                "sha256": str(source["sha256"]),
                "page_count": int(source["page_count"]),
            }
            for source_id, source in supplemental_sources.items()
        },
        "docx": {
            "filename": docx_path.name,
            "size_bytes": docx_path.stat().st_size,
            "sha256": sha256_file(docx_path),
            "expected_cover_pages": len(data["items"]),
            "source_cover_pages": [
                int(item["cover_source_page"]) for item in data["items"]
            ],
            "style_tokens": DOCX_STYLE_TOKENS,
        },
        "catalog_docx": {
            "filename": catalog_docx_path.name,
            "size_bytes": catalog_docx_path.stat().st_size,
            "sha256": sha256_file(catalog_docx_path),
            "signature_status": "待申请人签名或盖章",
        },
        "pdfs": records,
        "pdf_file_count": len(records),
        "evidence_page_count": sum(int(record["page_count"]) for record in records),
        "pagination_policy": data["page_number_policy"],
        "limits": {
            "max_pdf_files": MAX_PDF_FILES,
            "max_each_pdf_bytes": MAX_PDF_BYTES,
        },
    }
    audit_path.write_text(
        json.dumps(audit_data, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(
        json.dumps(
            {
                "docx": str(docx_path),
                "catalog_docx": str(catalog_docx_path),
                "pdf_count": len(records),
                "evidence_pages": audit_data["evidence_page_count"],
                "audit_report": str(audit_path),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except ValueError as error:
        print(f"错误：{error}", file=sys.stderr)
        raise SystemExit(2)
